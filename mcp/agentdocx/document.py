"""Document wrapper combining python-docx with direct OOXML access."""

from __future__ import annotations

import copy
import uuid
import zipfile
import shutil
import tempfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument_
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from lxml import etree

from .oxml_helpers import (
    W, NSMAP,
    tag, make_element, find, child, children,
    text_content, get_paragraph_runs, find_text_position,
    get_next_comment_id, get_next_revision_id,
    ensure_pPr, ensure_rPr,
)


# ── Data classes ────────────────────────────────────────────────


@dataclass
class ParaInfo:
    """Information about a paragraph in the document."""
    index: int
    text: str
    style: str = ""
    is_heading: bool = False
    heading_level: int = 0


@dataclass
class DocInfo:
    """Document metadata returned by get_info."""
    path: str
    paragraph_count: int
    paragraphs: list[ParaInfo] = field(default_factory=list)


# ── Document class ──────────────────────────────────────────────


class DocxDocument:
    """Wraps a python-docx Document with direct OOXML access for advanced features.

    Provides:
    - Open/save .docx files
    - Access to paragraphs and runs via python-docx API
    - Direct lxml element access for OOXML manipulation
    - Comments part management
    - Track changes support
    """

    def __init__(self, path: str | Path):
        self._path = Path(path)
        if not self._path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        self._docx = DocxDocument_(str(self._path))
        self._body: etree._Element = self._docx.element.body

        # Comments part
        self._comments: Optional[etree._Element] = None
        self._comments_part = None
        self._load_comments()

    # ── Properties ──────────────────────────────────────────

    @property
    def path(self) -> Path:
        return self._path

    @property
    def docx(self) -> DocxDocument_:
        return self._docx

    @property
    def body(self) -> etree._Element:
        return self._body

    @property
    def paragraphs(self):
        """Return paragraph elements directly."""
        return children(self._body, "p")

    @property
    def paragraph_count(self) -> int:
        return len(self.paragraphs)

    # ── Comments ────────────────────────────────────────────

    def _load_comments(self):
        """Load the comments part if it exists."""
        try:
            comments_part = self._docx.part.part_related_by(RT.COMMENTS)
        except (KeyError, ValueError):
            self._comments = None
            self._comments_part = None
            return

        if comments_part is not None:
            self._comments_part = comments_part
            self._comments = etree.fromstring(comments_part.blob)

    def _create_comments_element(self):
        """Create a fresh w:comments element (without adding to package yet)."""
        comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:comments xmlns:w="{NSMAP["w"]}"'
            f' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
            f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'</w:comments>'
        )
        self._comments = etree.fromstring(comments_xml.encode("utf-8"))

    @property
    def comments_element(self) -> Optional[etree._Element]:
        return self._comments

    def comments_count(self) -> int:
        if self._comments is None:
            return 0
        return len([c for c in self._comments if c.tag == tag("comment")])

    # ── Text access ─────────────────────────────────────────

    def get_paragraph_text(self, index: int) -> str:
        """Get the text of a paragraph by index."""
        paras = self.paragraphs
        if index < 0 or index >= len(paras):
            raise IndexError(f"Paragraph index {index} out of range (0-{len(paras) - 1})")
        return text_content(paras[index])

    def get_full_text(self) -> str:
        """Get the full text of the document with paragraph breaks."""
        texts = []
        for p in self.paragraphs:
            texts.append(text_content(p))
        return "\n".join(texts)

    def search_text(self, query: str) -> list[dict]:
        """Search for text in the document.

        Returns list of dicts with paragraph_index, offset, and context.
        """
        results = []
        for pi, para in enumerate(self.paragraphs):
            para_text = text_content(para)
            start = 0
            while True:
                idx = para_text.find(query, start)
                if idx == -1:
                    break
                # Get some context around the match
                ctx_start = max(0, idx - 20)
                ctx_end = min(len(para_text), idx + len(query) + 20)
                context = para_text[ctx_start:ctx_end]
                results.append({
                    "paragraph_index": pi,
                    "offset": idx,
                    "match_length": len(query),
                    "context": context,
                    "paragraph_text": para_text,
                })
                start = idx + 1
        return results

    def get_info(self) -> DocInfo:
        """Get document metadata and structure info."""
        paras_info = []
        for i, p in enumerate(self.paragraphs):
            txt = text_content(p)
            ppr = child(p, "pPr")
            style = ""
            is_heading = False
            heading_level = 0

            if ppr is not None:
                pstyle = child(ppr, "pStyle")
                if pstyle is not None:
                    style = pstyle.get(W + "val", "")

            if style and "Heading" in style or "heading" in style:
                is_heading = True
                # Extract level number
                for c in style:
                    if c.isdigit():
                        heading_level = int(c)
                        break

            paras_info.append(ParaInfo(
                index=i,
                text=txt[:200] + ("..." if len(txt) > 200 else ""),
                style=style,
                is_heading=is_heading,
                heading_level=heading_level,
            ))

        return DocInfo(
            path=str(self._path),
            paragraph_count=len(paras_info),
            paragraphs=paras_info,
        )

    # ── Save ────────────────────────────────────────────────

    def save(self, path: str | Path | None = None):
        """Save the document, optionally to a new path.

        Uses the original file as a ZIP base and injects only modified parts
        (document.xml and comments.xml). This preserves the original ZIP
        structure, directory entries, and unmodified XML files.
        """
        save_path = Path(path) if path else self._path

        # Serialize modified document and comments
        has_comments = (
            self._comments is not None
            and len([c for c in self._comments if c.tag == tag("comment")]) > 0
        )

        doc_xml = etree.tostring(
            self._docx.element, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        comments_xml = None
        if has_comments:
            comments_xml = etree.tostring(
                self._comments, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        # Build output ZIP from original structure
        self._write_package(save_path, doc_xml, comments_xml)

        if path:
            self._path = save_path
        return str(save_path)

    def save_to_bytes(self) -> bytes:
        """Save the document to bytes."""
        buf = BytesIO()
        # Use a temp file and read it back
        import tempfile, os
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = tmp.name
        tmp.close()
        try:
            self.save(tmp_path)
            with open(tmp_path, "rb") as f:
                buf.write(f.read())
        finally:
            os.unlink(tmp_path)
        return buf.getvalue()

    def _write_package(self, save_path: Path, doc_xml: bytes, comments_xml: bytes | None):
        """Write the output ZIP by injecting modified files into the original ZIP.

        This preserves directory entries, file ordering, and avoids re-serializing
        untouched XML files (styles, settings, etc.) which could introduce
        formatting differences (line endings, quote styles) that Office rejects.
        """
        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

        tmp_path = save_path.with_suffix(".tmp.docx")

        with zipfile.ZipFile(str(self._path), "r") as zin:
            with zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    fname = item.filename

                    # Skip existing comments — will write fresh if needed
                    if fname == "word/comments.xml" and comments_xml:
                        continue

                    data = zin.read(fname)

                    if fname == "word/document.xml":
                        data = doc_xml

                    elif fname == "[Content_Types].xml" and comments_xml:
                        ct = etree.fromstring(data)
                        if not any(
                            ov.get("PartName") == "/word/comments.xml"
                            for ov in ct.findall(f"{{{ct_ns}}}Override")
                        ):
                            ov = etree.SubElement(ct, f"{{{ct_ns}}}Override")
                            ov.set("PartName", "/word/comments.xml")
                            ov.set("ContentType",
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
                        data = etree.tostring(ct, xml_declaration=True, encoding="UTF-8", standalone=True)

                    elif fname == "word/_rels/document.xml.rels" and comments_xml:
                        rels_xml = etree.fromstring(data)
                        max_rid = max(
                            (int(r.get("Id", "rId0")[3:]) for r in rels_xml if r.get("Id", "").startswith("rId")),
                            default=0,
                        )
                        if not any(r.get("Type") == RT.COMMENTS for r in rels_xml):
                            cr = etree.SubElement(rels_xml, "Relationship")
                            cr.set("Id", f"rId{max_rid + 1}")
                            cr.set("Type", RT.COMMENTS)
                            cr.set("Target", "comments.xml")
                        data = etree.tostring(rels_xml, xml_declaration=True, encoding="UTF-8", standalone=True)
                        zout.writestr(item, data)
                        # Inject comments.xml right after rels
                        zout.writestr("word/comments.xml", comments_xml)
                        continue

                    zout.writestr(item, data)

                # Fallback: if comments weren't written (e.g., no rels file)
                if comments_xml and "word/comments.xml" not in zin.namelist():
                    # Already written above via the rels branch — this is a safety net
                    pass

        # Atomic replace
        if save_path.exists():
            save_path.unlink()
        shutil.move(str(tmp_path), str(save_path))

    # ── Paragraph operations ────────────────────────────────

    def add_paragraph(self, text: str = "", style: str | None = None) -> int:
        """Add a new paragraph to the end of the document.

        Returns the index of the new paragraph.
        """
        p = make_element("p")
        ppr = make_element("pPr")
        p.append(ppr)

        if style:
            from .oxml_helpers import set_or_replace
            set_or_replace(ppr, "pStyle", attrib={W + "val": style})

        if text:
            r = make_element("r")
            t = make_element("t", text=text)
            r.append(t)
            p.append(r)

        self._body.append(p)
        return len(self.paragraphs) - 1

    def insert_paragraph(self, index: int, text: str = "", style: str | None = None) -> int:
        """Insert a paragraph at the given index.

        Returns the index of the new paragraph.
        """
        paras = self.paragraphs
        if index < 0 or index > len(paras):
            raise IndexError(f"Paragraph index {index} out of range (0-{len(paras)})")

        p = make_element("p")
        ppr = make_element("pPr")
        p.append(ppr)

        if style:
            from .oxml_helpers import set_or_replace
            set_or_replace(ppr, "pStyle", attrib={W + "val": style})

        if text:
            r = make_element("r")
            t = make_element("t", text=text)
            r.append(t)
            p.append(r)

        if index == len(paras):
            self._body.append(p)
        else:
            target = paras[index]
            from .oxml_helpers import insert_before
            insert_before(target, p)

        return index

    def delete_paragraph(self, index: int):
        """Delete a paragraph by index."""
        paras = self.paragraphs
        if index < 0 or index >= len(paras):
            raise IndexError(f"Paragraph index {index} out of range (0-{len(paras) - 1})")
        self._body.remove(paras[index])
